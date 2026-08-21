import hashlib,zipfile
import pytest
from docx import Document
from cko_local_finder.domain.models import DiscoveredFile,ExtractionPolicy
from cko_local_finder.infrastructure.extractors import DocxExtractor,ExtractionError

def source(path):
    b=path.read_bytes(); d=hashlib.sha256(b).hexdigest(); return DiscoveredFile(d,str(path),path.name,d,len(b),".docx",1,"application/vnd.openxmlformats-officedocument.wordprocessingml.document")

def test_paragraph_table_order_and_determinism(tmp_path):
    path=tmp_path/"a.docx"; doc=Document(); doc.add_paragraph("Before"); table=doc.add_table(rows=1,cols=2); table.cell(0,0).text="A"; table.cell(0,1).text="B"; doc.add_paragraph("After"); doc.save(path)
    result=DocxExtractor().extract(source(path)); assert result.text=="Before\nA\tB\nAfter"; assert result==DocxExtractor().extract(source(path))

def test_empty_and_corrupt_docx(tmp_path):
    empty=tmp_path/"empty.docx"; Document().save(empty); assert DocxExtractor().extract(source(empty)).status=="EMPTY"
    bad=tmp_path/"bad.docx"; bad.write_bytes(b"PKbad")
    with pytest.raises(ExtractionError) as error: DocxExtractor().extract(source(bad))
    assert error.value.code=="CORRUPT_DOCX"

@pytest.mark.parametrize(("name","policy","code"),[("../evil",ExtractionPolicy(),"UNSAFE_DOCX_PATH"),("extra",ExtractionPolicy(max_docx_archive_entries=1),"DOCX_TOO_MANY_ENTRIES"),("extra",ExtractionPolicy(max_docx_uncompressed_bytes=1),"DOCX_UNCOMPRESSED_TOO_LARGE")])
def test_zip_safety(tmp_path,name,policy,code):
    path=tmp_path/"unsafe.docx"
    with zipfile.ZipFile(path,"w") as z:
        z.writestr("[Content_Types].xml","x"); z.writestr("word/document.xml","x"); z.writestr(name,"x")
    with pytest.raises(ExtractionError) as error: DocxExtractor(policy).extract(source(path))
    assert error.value.code==code
