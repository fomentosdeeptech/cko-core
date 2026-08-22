from __future__ import annotations

import hashlib
from io import StringIO
import json
from pathlib import Path
import sqlite3

from docx import Document

from cko_local_finder.cli.main import main
from cko_local_finder.infrastructure.sqlite import SQLiteDocumentRepository
from tests.corpus_factory import materialize_corpus


def invoke(*args: str) -> tuple[int, str, str]:
    stdout, stderr = StringIO(), StringIO()
    code = main(args, stdout=stdout, stderr=stderr)
    return code, stdout.getvalue(), stderr.getvalue()


def richer_corpus(root: Path) -> dict[str, str]:
    materialize_corpus(root, size_limit=128)
    (root / "edge" / "invalid.txt").write_bytes(b"\xff\xfe\x00")
    (root / "valid" / "diacritics.txt").write_text(
        "governança inovação café evidência pesquisável", encoding="utf-8"
    )
    document = Document()
    document.add_paragraph("governança sintética em parágrafo")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "controle"
    table.cell(0, 1).text = "local"
    table.cell(1, 0).text = "evidência"
    table.cell(1, 1).text = "determinística"
    document.save(root / "valid" / "paragraphs-and-table.docx")
    return {
        path.relative_to(root).as_posix(): hashlib.sha256(path.read_bytes()).hexdigest()
        for path in sorted(root.rglob("*")) if path.is_file()
    }


def test_complete_workflow_idempotency_filters_provenance_and_reports(tmp_path: Path) -> None:
    root, database = tmp_path / "corpus", tmp_path / "finder.sqlite"
    source_hashes = richer_corpus(root)
    first = invoke("ingest", str(root), "--database", str(database), "--format", "json")
    second = invoke("ingest", str(root), "--database", str(database), "--format", "json")
    assert first[0] == second[0] == 1 and not first[2] and not second[2]
    first_payload, second_payload = json.loads(first[1]), json.loads(second[1])
    assert first_payload["discovered_documents"] == second_payload["discovered_documents"] == 14
    assert first_payload["unique_documents"] == second_payload["unique_documents"] == 13
    assert first_payload["locations"] == second_payload["locations"] == 14
    assert first_payload["duplicate_groups"] == second_payload["duplicate_groups"] == 1
    assert first_payload["recoverable_failures"] == second_payload["recoverable_failures"] == 3

    base = ("search", "café", "--database", str(database), "--format", "json")
    code, stdout, stderr = invoke(*base)
    result = json.loads(stdout)
    assert code == 0 and not stderr and result["total_matches"] == 1
    hit = result["results"][0]
    assert "[[café]]" in hit["snippet"] and hit["path"] == "valid/diacritics.txt"
    digest = hit["sha256"]
    filters = (
        ("--extension", "txt"), ("--media-type", "text/plain"),
        ("--root", str(root.resolve())), ("--path-prefix", "valid/"),
        ("--sha256", digest),
    )
    for option in filters:
        assert json.loads(invoke(*base, *option)[1])["total_matches"] == 1
    combined = sum((list(option) for option in filters), [])
    assert json.loads(invoke(*base, *combined)[1])["total_matches"] == 1
    assert invoke(*base, "--path-prefix", "../escape")[0] == 2

    shown = json.loads(invoke("show", digest, "--database", str(database), "--format", "json")[1])
    document = shown["document"]
    assert document["sha256"] == digest and len(document["origins"]) == 1
    assert document["extraction"]["status"] == "SUCCESS" and document["indexing"]["indexed"] is True
    duplicate_report = json.loads(invoke("duplicates", "--database", str(database), "--format", "json")[1])
    assert len(duplicate_report["duplicates"]) == 1
    assert len(duplicate_report["duplicates"][0]["origins"]) == 2

    for report_type in ("ingestion", "failures", "duplicates"):
        args = ["report", report_type, "--database", str(database), "--format", "json"]
        if report_type == "ingestion":
            args.extend(("--root", str(root.resolve())))
        code, output, error = invoke(*args)
        assert code == 0 and not error and output.endswith("\n") and not output.endswith("\n\n")
        assert json.loads(output) is not None and "traceback" not in output.lower() and "\x1b[" not in output

    assert source_hashes == richer_hashes(root)


def richer_hashes(root: Path) -> dict[str, str]:
    return {
        path.relative_to(root).as_posix(): hashlib.sha256(path.read_bytes()).hexdigest()
        for path in sorted(root.rglob("*")) if path.is_file()
    }


def test_rebuild_from_sources_is_semantically_equivalent(tmp_path: Path) -> None:
    root = tmp_path / "corpus"
    richer_corpus(root)
    snapshots = []
    for name in ("first.sqlite", "rebuilt.sqlite"):
        database = tmp_path / name
        assert invoke("ingest", str(root), "--database", str(database), "--format", "json")[0] == 1
        repository = SQLiteDocumentRepository(database)
        with repository.connection() as connection:
            tables = {}
            for table, columns, order in (
                ("documents", "sha256,size_bytes,extension,media_type,physical_metadata_json", "sha256"),
                ("document_locations", "document_sha256,root,relative_path,observed_size_bytes,mtime_ns", "root,relative_path"),
                ("extractions", "document_sha256,extractor,extractor_version,status,text_content,metadata_json", "document_sha256"),
                ("processing_issues", "document_sha256,root,relative_path,stage,code,message,recoverable", "stage,relative_path,code"),
                ("search_index_documents", "document_sha256,title,body,extension,media_type,root,relative_path", "document_sha256"),
            ):
                rows = connection.execute(f"SELECT {columns} FROM {table} ORDER BY {order}").fetchall()
                tables[table] = [tuple(row) for row in rows]
        search = json.loads(invoke("search", "governança", "--database", str(database), "--format", "json")[1])
        duplicates = json.loads(invoke("duplicates", "--database", str(database), "--format", "json")[1])
        snapshots.append((tables, search["results"], duplicates["duplicates"]))
    assert snapshots[0] == snapshots[1]


def test_same_relative_path_under_two_roots_keeps_distinct_provenance(tmp_path: Path) -> None:
    database = tmp_path / "finder.sqlite"
    roots = (tmp_path / "root-a", tmp_path / "root-b")
    for index, root in enumerate(roots):
        (root / "same").mkdir(parents=True)
        (root / "same" / "note.txt").write_text(f"root identity {index}", encoding="utf-8")
        assert invoke("ingest", str(root), "--database", str(database))[0] == 0
    repository = SQLiteDocumentRepository(database)
    first = repository.provenance_by_location(str(roots[0].resolve()), "same/note.txt")
    second = repository.provenance_by_location(str(roots[1].resolve()), "same/note.txt")
    assert first is not None and second is not None
    assert first.document.sha256 != second.document.sha256
